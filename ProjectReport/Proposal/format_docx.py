#!/usr/bin/env python3
"""Post-process pandoc docx: add table borders, apply heading styles, set fonts."""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_table_borders(table, color="000000", size="4"):
    """Add all borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def shade_header_row(table, color="D9E2F3"):
    """Shade the first row of a table."""
    for cell in table.rows[0].cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def style_headings(doc):
    """Apply consistent heading formatting."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # dark navy
                if level == '1':
                    run.font.size = Pt(22)
                elif level == '2':
                    run.font.size = Pt(16)
                elif level == '3':
                    run.font.size = Pt(13)


def set_body_font(doc, font_name="Calibri", font_size=Pt(11)):
    """Set default body font."""
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('Heading'):
            for run in paragraph.runs:
                run.font.name = font_name
                if run.font.size is None:
                    run.font.size = font_size


def set_table_font(doc, font_name="Calibri", font_size=Pt(9)):
    """Set table cell font."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size


def center_images(doc):
    """Centre all paragraphs that contain images."""
    for paragraph in doc.paragraphs:
        if paragraph._element.findall(f'.//{qn("w:drawing")}') or \
           paragraph._element.findall(f'.//{qn("w:pict")}'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_breaks(doc):
    """Add page break before each Heading 2 (## sections), except the first."""
    seen_first_h2 = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 2':
            if not seen_first_h2:
                seen_first_h2 = True
                continue
            # Set "page break before" paragraph property
            pPr = paragraph._element.get_or_add_pPr()
            page_break = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="true"/>')
            pPr.append(page_break)


def add_page_numbers(doc):
    """Add 'Page X of Y' footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # "Page "
        run1 = p.add_run("Page ")
        run1.font.size = Pt(9)
        run1.font.name = "Calibri"
        # PAGE field
        fld_page = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        run1._element.addnext(fld_page)
        # " of "
        run2_xml = parse_xml(
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr><w:t xml:space="preserve"> of </w:t></w:r>'
        )
        fld_page.addnext(run2_xml)
        # NUMPAGES field
        fld_total = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        run2_xml.addnext(fld_total)


def set_narrow_margins(doc):
    """Set narrow page margins."""
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def main():
    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} <file.docx>")
        sys.exit(1)

    path = sys.argv[1]
    doc = Document(path)

    # Page setup
    set_narrow_margins(doc)

    # Page numbers
    add_page_numbers(doc)

    # Headings
    style_headings(doc)

    # Centre images
    center_images(doc)

    # Body font
    set_body_font(doc)

    # Tables: borders, header shading, font
    for table in doc.tables:
        set_table_borders(table)
        shade_header_row(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_font(doc)

    doc.save(path)
    print(f"Formatted: {path}")


if __name__ == "__main__":
    main()
